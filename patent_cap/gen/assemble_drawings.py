"""Assemble Drawings_ACN1408_CAP.docx — one figure per labelled A4 sheet.

Reconstructs the sheet header per page (Applicant / Total no. of sheets / Sheet no /
Figure no), reusing the field labels of the provisional's Drawings_ACN1408.docx (which
packs two figures on a single sheet — so this is a reconstruction, not a copy). Embeds the
1-bit PNG of each figure.
"""
import pathlib

from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# Each figure must fit ON ITS OWN A4 sheet without overflowing (the cause of the
# "figures pile on each other" issue): constrain by BOTH width and height.
MAX_W_IN = 6.4
MAX_H_IN = 7.8


def fit_width(png_path):
    """Return the display width (Inches) that fits the image within MAX_W x MAX_H."""
    w, h = Image.open(png_path).size
    ar = w / h
    disp_w = MAX_W_IN
    if disp_w / ar > MAX_H_IN:      # too tall -> constrain by height
        disp_w = MAX_H_IN * ar
    return Inches(disp_w)


HERE = pathlib.Path(__file__).resolve().parent
ROOT = HERE.parent
FIG_DIR = ROOT / "mermaid"   # Mermaid (library) renders — professional, not hand-drawn

APPLICANT = "Amity University"
SHEETS = [
    (1, "fig1_system_architecture", "System architecture"),
    (2, "fig2_sequence", "Sequence of system operation"),
    (3, "fig3_ai_framework", "AI framework / signal-processing pipeline"),
    (4, "fig4_model_architecture", "BiLSTM + temporal-attention model architecture"),
    (5, "fig5_vehicle_state_machine", "Fail-safe vehicle-control state machine"),
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)  # A4 portrait
    sec.left_margin = sec.right_margin = Mm(20)
    sec.top_margin = sec.bottom_margin = Mm(18)

    total = len(SHEETS)
    for i, (num, fname, title) in enumerate(SHEETS):
        if i > 0:
            doc.add_page_break()

        header = doc.add_paragraph()
        run = header.add_run(
            f"Applicant: {APPLICANT}        Total no. of sheets: {total:02d}        "
            f"Sheet no: {num:02d}")
        run.font.size = Pt(10)

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(f"FIG. {num} — {title}")
        crun.bold = True
        crun.font.size = Pt(11)

        png = FIG_DIR / f"{fname}.png"
        img = doc.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.add_run().add_picture(str(png), width=fit_width(png))

    # closing applicant block on the last sheet
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(f"\n{APPLICANT}\nName of Applicant\nSignature:").font.size = Pt(10)

    out = ROOT / "Drawings_ACN1408_CAP.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    p = build()
    print("wrote", p)
