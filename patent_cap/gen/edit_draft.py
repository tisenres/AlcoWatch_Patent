"""Edit the ACN1408 provisional draft for the CAP:
  - apply the reconciliation corrections (code-vs-provisional) IN RED;
  - insert the AI framework / model EXPLANATION and Figure 2-5 descriptions IN RED.
Every change/addition is coloured red so the supervisor sees exactly what we touched.
Output: a new file; the original is left intact.
"""
import sys
import pathlib

from docx import Document
from docx.shared import RGBColor, Pt

RED = RGBColor(0xC0, 0x00, 0x00)

SRC = pathlib.Path("ACN1408 Prov._Patent_130625 (1).docx")
OUT = pathlib.Path("patent_cap/ACN1408_CAP_Draft_REVISED.docx")

# (old phrase -> new phrase). New text is rendered red; the rest of the paragraph stays black.
REPLACEMENTS = [
    ("transdermal and physiological sensors—specifically photoplethysmography (PPG), "
     "electrodermal activity (EDA), and skin temperature sensors",
     "physiological sensors—specifically photoplethysmography (PPG), an electrodermal-"
     "activity (EDA) channel estimated from heart-rate variability, and a skin-temperature sensor"),
    ("transdermal and physiological sensors—such as PPG, EDA, and skin temperature sensors",
     "physiological sensors—such as PPG, an estimated electrodermal-activity (EDA) channel, "
     "and skin temperature sensors"),
    ("electrodermal activity (EDA) for quantifying sympathetic nervous system responses",
     "electrodermal activity (EDA), estimated/derived from heart-rate variability, for "
     "quantifying sympathetic nervous system responses"),
    ("interpreting the raw sensor data in real time to infer transdermal alcohol "
     "concentrations, which are then computationally converted to estimated BAC values",
     "interpreting the raw sensor data in real time to infer the driver's blood alcohol "
     "concentration (BAC) from the fused physiological signals"),
    ("These algorithms are trained on large-scale physiological datasets",
     "These algorithms are trained on physiologically-modeled datasets (synthetic data in "
     "the present embodiment)"),
    ("The AI engine dynamically adjusts its BAC thresholds based on prior readings and "
     "environmental feedback",
     "The AI engine applies a fixed statutory BAC threshold (0.08 g/dL), while region-"
     "specific calibration adapts the BAC estimate based on prior readings and environmental feedback"),
    ("further fortified through encryption standards",
     "further fortified through AES-256 encryption (per specification)"),
]

# Figure descriptions added to BRIEF DESCRIPTION OF THE DRAWINGS (red).
FIGURE_LINES = [
    "Figure 2 illustrates a sequence diagram of the system operation, in accordance with "
    "embodiments of the present invention.",
    "Figure 3 illustrates the AI framework and signal-processing pipeline, in accordance "
    "with embodiments of the present invention.",
    "Figure 4 illustrates the neural-network model architecture (a bidirectional LSTM with "
    "a temporal attention mechanism), in accordance with embodiments of the present invention.",
    "Figure 5 illustrates the fail-safe vehicle-control state machine, in accordance with "
    "embodiments of the present invention.",
]

# AI framework / model EXPLANATION block (red), inserted into the Detailed Description.
EXPLANATION = [
    ("AI Framework and Model Architecture", True),  # (text, is_heading)
    ("Referring to Figures 3 and 4, the AI module (120) carried by the smartwatch (110) "
     "estimates blood alcohol concentration (BAC) through a multi-stage pipeline. Raw signals "
     "from the photoplethysmography sensor (112), the estimated electrodermal-activity channel "
     "(114), and the temperature sensor (116) are preprocessed and segmented into fixed-length "
     "temporal windows of ten time steps (122), forming a feature tensor (124) of shape ten time "
     "steps by six features (a PPG value, a PPG signal-quality measure, an EDA value, a skin "
     "temperature, an ambient temperature, and a humidity value).", False),
    ("As shown in Figure 4, the feature tensor (124) is processed by a bidirectional long "
     "short-term memory (BiLSTM) stage (126) of sixty-four units, producing a temporally-encoded "
     "representation, followed by a dropout stage. A temporal attention mechanism (128) computes "
     "per-time-step scores by a dense projection with a hyperbolic-tangent activation, normalizes "
     "them by a softmax to obtain attention weights over the ten time steps, and forms a weighted "
     "sum across time to produce an attention context vector. The context vector is passed through "
     "a dense regression head (130) comprising a thirty-two-unit rectified-linear layer, a dropout "
     "stage, a sixteen-unit rectified-linear layer, and a final single-unit linear layer that "
     "outputs the BAC estimate in grams per decilitre.", False),
    ("The model is trained with a safety-prioritized asymmetric loss function in which "
     "false-negative errors (predicting a sober reading when the true BAC is above the threshold) "
     "are penalized thirty times more heavily than false-positive errors, thereby biasing the "
     "system toward the suppression of missed-intoxication events. The trained model is quantized "
     "and executed on the smartwatch by an on-device inference engine (118).", False),
    ("Referring to Figure 3, a climate-adaptive calibration stage (132) applies a deterministic "
     "post-inference correction to the model's raw BAC output using region-specific temperature and "
     "humidity coefficients; the calibration is a correction applied to the model output and is not "
     "a layer of the trained neural network. The calibrated BAC is compared by the decision logic "
     "(154) against a statutory threshold of 0.08 g/dL, and an allow/block determination, together "
     "with a confidence value, is transmitted over the secure BLE link (140) to the vehicle control "
     "module (150).", False),
    ("Referring to Figure 5, the decision logic (154) implements a fail-safe state machine whose "
     "default state upon power-on is ignition-blocked. Ignition is enabled only while the link is "
     "active, a recent BAC update has been received, the estimated BAC is below 0.08 g/dL, and the "
     "watch is worn; loss of communication beyond a timeout, removal of the watch, or stale data "
     "returns the system to the ignition-blocked state. A manual override (180), engaged by a "
     "sustained button hold, transitions to a time-limited, logged override state that "
     "automatically reverts to the ignition-blocked state.", False),
]


def rebuild_with_red(p, replacements):
    """Rebuild paragraph p: unchanged text stays black, replaced segments become red."""
    text = p.text
    spans = []
    for old, new in replacements:
        idx = text.find(old)
        if idx >= 0:
            spans.append((idx, idx + len(old), new))
    if not spans:
        return False
    spans.sort()
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    cur = 0
    for s, e, new in spans:
        if s > cur:
            p.add_run(text[cur:s])
        run = p.add_run(new)
        run.font.color.rgb = RED
        cur = e
    if cur < len(text):
        p.add_run(text[cur:])
    return True


def red_para_before(anchor, text, heading=False):
    p = anchor.insert_paragraph_before()
    run = p.add_run(text)
    run.font.color.rgb = RED
    if heading:
        run.bold = True
        run.font.size = Pt(12)
    return p


def main():
    doc = Document(str(SRC))
    n_edits = 0
    for p in doc.paragraphs:
        if rebuild_with_red(p, REPLACEMENTS):
            n_edits += 1

    # anchors
    anchor_detailed = None   # "DETAILED DESCRIPTION OF THE INVENTION"
    anchor_closing = None    # "It is to be understood that the above description..."
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("DETAILED DESCRIPTION OF THE INVENTION") and anchor_detailed is None:
            anchor_detailed = p
        if t.startswith("It is to be understood that the above description") and anchor_closing is None:
            anchor_closing = p

    # Figure 2-5 descriptions: before the DETAILED DESCRIPTION heading (i.e., end of Brief Description)
    if anchor_detailed is not None:
        for line in FIGURE_LINES:
            red_para_before(anchor_detailed, line)

    # Explanation block: before the closing boilerplate
    if anchor_closing is not None:
        for text, is_head in EXPLANATION:
            red_para_before(anchor_closing, text, heading=is_head)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"reconciliation paragraphs edited: {n_edits}")
    print(f"figure lines added: {len(FIGURE_LINES)} | explanation paragraphs added: {len(EXPLANATION)}")
    print("wrote", OUT)


if __name__ == "__main__":
    main()
