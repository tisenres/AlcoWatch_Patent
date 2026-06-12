"""Minimal Markdown -> .docx converter for the ACN1408 hand-off documents.

Handles the subset used by disclosure_ai_section.md and provisional_reconciliation.md:
ATX headings (#/##/###), paragraphs with **bold** and `code`, '- ' bullet lists,
GitHub-style pipe tables, blockquotes ('> '), and '---' rules. Not a general converter.

Usage:  python3 gen/md_to_docx.py <input.md> <output.docx> ["Doc Title"]
"""
import re
import sys
import pathlib

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_inline(paragraph, text):
    """Render **bold**, *italic*, and `code` spans into runs."""
    for tok in re.split(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Consolas"
        elif tok.startswith("*") and tok.endswith("*"):
            paragraph.add_run(tok[1:-1]).italic = True
        else:
            paragraph.add_run(tok)


def convert(md_path, docx_path, title=None):
    lines = pathlib.Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # tables: a header row followed by a |---| separator
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in block]
            rows = [rows[0]] + rows[2:]  # drop the separator row
            table = doc.add_table(rows=0, cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                cells = table.add_row().cells
                for ci, val in enumerate(row[:len(cells)]):
                    cells[ci].paragraphs[0].text = ""
                    add_inline(cells[ci].paragraphs[0], val)
                    if ri == 0:
                        for run in cells[ci].paragraphs[0].runs:
                            run.bold = True
            doc.add_paragraph()
            continue

        if not line:
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.strip() == "---":
            pass
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif line.startswith("> "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:]); r.italic = True
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    out = convert(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
    print("wrote", out)
